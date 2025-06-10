import streamlit as st
from datetime import date
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import re

# Configure Gemini
genai.configure(api_key="AIzaSyDtAUIRtLS1T1H5jQ2ZmcafeCenOB3vQAA")

# Streamlit UI Setup
st.set_page_config(page_title="Professional Daily Report Generator", layout="centered")
st.title("📋 Daily Report Generator")

# Form Inputs
with st.form("report_form"):
    name = st.text_input("👤 Name")
    report_date = st.date_input("📅 Date", value=date.today())
    morning_session = st.text_area("🌅 Morning / Forenoon Session")
    afternoon_session = st.text_area("🌇 Afternoon Session")
    completed_tasks = st.text_area("✅ Completed Tasks")
    pending_tasks = st.text_area("⏳ Pending Tasks")
    submitted = st.form_submit_button("Generate Report")

# Utility: Strip bullet points and asterisks
def clean_text(text):
    text = re.sub(r"[*•\-]+", "", text)
    return text.strip()

# Create styled paragraph
def add_bold_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

if submitted:
    # Prompt construction
    prompt = f"""
    Please write a detailed and formal daily work report based on the inputs below.
    Do not use bullet points or markdown. Expand the points into full, professional sentences.
    Use clean, formal language.

    Name: {name}
    Date: {report_date}
    Morning Session: {morning_session}
    Afternoon Session: {afternoon_session}
    Completed Tasks: {completed_tasks}
    Pending Tasks: {pending_tasks}
    """

    try:
        # Gemini call
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        report_text = clean_text(response.text)

        # Create .docx
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        doc.add_heading("Daily Work Report", level=1)

        add_bold_heading(doc, f"Name: {name}")
        add_bold_heading(doc, f"Date: {report_date.strftime('%B %d, %Y')}")
        doc.add_paragraph(report_text)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Report generated successfully!")

        st.download_button(
            label="📥 Download Report (.docx)",
            data=buffer,
            file_name=f"Daily_Report_{report_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"❌ Failed to generate report: {e}")
